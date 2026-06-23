# -*- coding: utf-8 -*-
"""Generate the IEEE CS Pro 2026 Event Concept & Execution Plan as a Google-Docs-friendly .docx."""
import shutil
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x15, 0x1B, 0x2E)
GOLD = RGBColor(0xB0, 0x88, 0x2C)
GREY = RGBColor(0x55, 0x5B, 0x66)

doc = Document()

# Base font
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x20, 0x24, 0x2C)


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=10.5, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color is not None:
        run.font.color.rgb = color


def h1(text):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = NAVY
    # gold bottom border
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "B0882C")
    pbdr.append(bottom)
    pPr.append(pbdr)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    return p


def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = GOLD
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    return p


def h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p


def body(text, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        set_cell_text(hdr[i], htext, bold=True, white=True, size=10.5)
        shade(hdr[i], "151B2E")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


# ---------------- TITLE BLOCK ----------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("IEEE CS Pro 2026")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Event Concept & Execution Plan")
r.font.size = Pt(13)
r.font.color.rgb = GOLD
r.bold = True

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tag.add_run("An exclusive leadership & networking forum for CS Professionals.")
r.italic = True
r.font.size = Pt(10.5)
r.font.color.rgb = GREY

# Quick facts table
facts = doc.add_table(rows=0, cols=2)
facts.style = "Table Grid"
fact_rows = [
    ("Event Theme", "CS Pro Conclave 2026 — Leaders. Learners. Impact."),
    ("Date", "Sunday, 26 July 2026"),
    ("Time", "9:00 AM – 1:00 PM (followed by lunch)"),
    ("Venue", "TBD"),
    ("Access", "By Invitation Only / Invite-based registration"),
    ("Organised by", "IEEE Computer Society – Bangalore Chapter"),
    ("In association with", "IEEE Bangalore Section · IEEE Computer Society 80 Years Celebration"),
    ("Contact", "ieeecspro@gmail.com  ·  +91 96089 53402"),
    ("Web / Social", "ieeecsbangalore.org  ·  Instagram @ieeecsbc"),
]
for k, v in fact_rows:
    cells = facts.add_row().cells
    set_cell_text(cells[0], k, bold=True, color=NAVY, size=10)
    shade(cells[0], "F3EFE2")
    set_cell_text(cells[1], v, size=10)
    cells[0].width = Inches(1.9)
    cells[1].width = Inches(4.6)

# ---------------- 1. OVERVIEW ----------------
h1("1.  Event Overview")
body("IEEE CS Pro 2026 is an exclusive leadership and networking forum designed to bring together leaders from "
     "industry, startups, academia and the broader innovation ecosystem to discuss emerging challenges, share best "
     "practices, and foster meaningful collaborations that shape the future workforce.")
body("The event is centred around a high-impact Leadership Round Table Conference, where leaders engage in focused "
     "discussions aimed at generating practical insights, actionable recommendations and long-term collaborative "
     "opportunities.")
body("The programme then expands into CS Professional Connect, a networking forum enabling leaders, professionals, "
     "HRs, academicians, Training & Placement Officers (TPOs) and IEEE members to interact and build lasting "
     "connections.")

# ---------------- 2. OBJECTIVES ----------------
h1("2.  Event Objectives")
for item in [
    "Facilitate meaningful dialogue among leaders from industry, startups and academia.",
    "Generate actionable insights on workforce readiness and innovation.",
    "Promote collaboration between institutions and employers.",
    "Encourage IEEE membership and community engagement.",
    "Build a sustainable ecosystem for future industry–academia partnerships.",
]:
    bullet(item)

# ---------------- 3. PARTICIPANT PROFILE ----------------
h1("3.  Participant Profile")
h2("Leadership Round Table (Leaders)")
make_table(
    ["Participants", "Expected strength"],
    [
        ["Industry Leaders · Startup Founders · Entrepreneurs · CXOs and Decision Makers · Invited Guests", "20–25"],
    ],
    widths=[5.0, 1.5],
)
h2("CS Professional Connect (Professionals)")
make_table(
    ["Participants", "Expected strength"],
    [
        ["HR Professionals · TPOs · Placement Coordinators · Academicians · Industry Professionals & Corporates · "
         "IEEE Members", "40–50"],
    ],
    widths=[5.0, 1.5],
)
h2("Organising Team")
make_table(
    ["Team", "Approx. number"],
    [
        ["IEEE ExCom Members", "~10"],
        ["Volunteers", "~10"],
    ],
    widths=[5.0, 1.5],
)

# ---------------- 4. RUN OF SHOW ----------------
h1("4.  Programme Flow / Run of Show")
body("Indicative timeline for the day (tentative — to be confirmed closer to the event):", italic=True, color=GREY)
make_table(
    ["Time", "Segment"],
    [
        ["9:00 AM onwards", "Arrival, registration, IEEE Membership Desk, refreshments & informal networking"],
        ["9:30 – 10:00 AM", "Inaugural Session"],
        ["10:00 – 10:45 AM", "Leadership Round Table Discussions (3 parallel tables)"],
        ["10:45 – 11:05 AM", "Rapid Fire Leadership Exchange"],
        ["11:05 – 11:20 AM", "Leadership Insights Presentation (one leader per table)"],
        ["11:20 – 11:30 AM", "Moderator Summary & Group Photograph"],
        ["11:30 – 11:50 AM", "Tea Break & Transition · Professional entry · IEEE Membership Drive"],
        ["11:50 – 12:00 PM", "Opening Address (Session 2)"],
        ["12:00 – 12:10 PM", "Ice Breakers & Introductions"],
        ["12:10 – 12:50 PM", "Meet & Greet Networking Session"],
        ["12:50 – 1:00 PM", "Closing Remarks & Acknowledgements"],
        ["1:00 PM onwards", "Lunch & Informal Networking"],
        ["Post-Lunch", "IEEE Computer Society 80 Years Celebration"],
    ],
    widths=[1.7, 4.8],
)

# ---------------- 5. SESSION 1 ----------------
h1("5.  Session 1 — Leadership Round Table Conference")

h2("Arrival & Welcome — 9:00 AM onwards")
for item in [
    "Registration Desk open.",
    "IEEE Membership Desk open.",
    "Welcome refreshments and snacks for leaders.",
    "Informal networking.",
    "Distribution of ID cards and event kits.",
]:
    bullet(item)

h2("Inaugural Session — 9:30 AM to 10:00 AM")
make_table(
    ["Time", "Programme"],
    [
        ["9:30 – 9:45 AM", "Inauguration · Welcome Address · Introduction to IEEE CS Pro"],
        ["9:45 – 9:55 AM", "Remarks by IEEE Dignitaries / ExCom Members"],
        ["9:55 – 10:00 AM", "Round Table Briefing · Discussion Guidelines · Table Allocation Instructions"],
    ],
    widths=[1.7, 4.8],
)

h2("Leadership Round Table Conference — 10:00 AM to 11:30 AM")
h3("Seating Format")
for item in [
    "Leaders seated at dedicated round tables positioned at the front.",
    "Goodies presented warmly at each table.",
    "One volunteer assigned to every table for coordination and documentation.",
    "IEEE ExCom members distributed across tables as facilitators.",
]:
    bullet(item)

h3("Round Table Discussions")
body("The leadership session consists of three parallel round table discussions.")

h3("Round Table 1 — Driving Organisational Transformation")
body("Actions Taken, Changes Implemented & Outcomes Achieved.", italic=True, color=GREY)
body("Discussion focus:")
for item in [
    "What significant changes are currently taking place within organisations?",
    "What initiatives have already been implemented?",
    "What strategic actions have been taken?",
    "What measurable outcomes, lessons and insights have emerged?",
]:
    bullet(item)
body("Expected outcome: ", italic=True)
bullet("Practical industry insights and real-world case studies that can inspire transformation across sectors.")

h3("Round Table 2 — Building the AI-Ready Workforce")
body("How Startups & Industry Can Build the AI-Ready Workforce Together.", italic=True, color=GREY)
body("Discussion focus:")
for item in [
    "Academia–industry collaboration.",
    "Internships and experiential learning.",
    "Curriculum modernisation.",
    "Future hiring expectations.",
    "Industry participation in workforce development.",
]:
    bullet(item)
body("Expected outcome: ", italic=True)
bullet("Actionable recommendations for institutions, employers and ecosystem stakeholders.")

h3("Round Table 3 — Beyond the Machine")
body("Building Human-Centred, Ethical & Purpose-Driven Innovation.", italic=True, color=GREY)
body("Discussion focus:")
for item in [
    "What differentiates great organisations beyond technological advancement?",
    "How can soul, responsibility and human values be embedded into innovation?",
    "What role do leadership, culture, trust and purpose play in an AI-driven world?",
    "How can organisations ensure technology serves humanity rather than merely advancing capability?",
    "What uniquely human qualities should be preserved and strengthened as AI becomes more capable?",
]:
    bullet(item)
body("Expected outcome: ", italic=True)
bullet("A deeper understanding of the human, ethical and purpose-driven dimensions required to thrive in an "
       "AI-powered future.")

h2("Leadership Insights & Knowledge Sharing")
h3("10:45 – 11:05 AM · Rapid Fire Leadership Exchange")
for item in [
    "Moderated rapid-fire discussion involving all leaders.",
    "Reflections from each table.",
    "Emerging trends, opportunities and challenges.",
]:
    bullet(item)
h3("11:05 – 11:20 AM · Leadership Insights Presentation")
body("Organisers nominate one representative leader from each table to present key insights, recommendations and "
     "takeaways. Format: 3 leaders, 5 minutes each.")
h3("11:20 – 11:30 AM · Moderator Summary & Group Photograph")
for item in [
    "Consolidated event insights.",
    "Group photograph.",
    "Transition to the networking session.",
]:
    bullet(item)

# ---------------- 6. TEA BREAK ----------------
h1("6.  Tea Break & Transition — 11:30 AM to 11:50 AM")
for item in [
    "Tea and refreshments.",
    "Professional entry and registrations.",
    "IEEE Membership Drive.",
    "Informal interactions between leaders and attendees.",
]:
    bullet(item)

# ---------------- 7. SESSION 2 ----------------
h1("7.  Session 2 — CS Professional Connect")
body("Time: 11:50 AM – 1:00 PM", color=GREY)
h2("Participants")
body("Professionals · HRs · TPOs · Placement Coordinators · Academicians · Industry Leaders · IEEE Members")
h2("Seating Arrangement")
bullet("Leaders remain seated at designated front tables.")
bullet("Professionals seated in an audience-style arrangement.")

h2("Opening Address — 11:50 AM to 12:00 PM")
body("Delivered by an IEEE ExCom Member.")
body("Theme: Navigating a Changing World Together — Responding to Emerging Global and Technological Challenges across "
     "industry, society and the geopolitical landscape.", italic=True)
body("Focus areas:")
for item in [
    "Rapid technological transformation.",
    "AI-driven disruption.",
    "Workforce readiness.",
    "Responsible innovation.",
    "Importance of collaboration across sectors.",
    "Current global developments and their impact on industries, institutions and society.",
    "How industry, academia and professional communities can work together to address shared challenges.",
]:
    bullet(item)

h2("Engagement & Networking Activities")
h3("12:00 – 12:10 PM · Ice Breakers & Introductions")
for item in [
    "Structured networking activities.",
    "Participant introductions.",
    "Cross-sector engagement.",
]:
    bullet(item)
h3("12:10 – 12:50 PM · Meet & Greet Networking Session")
for item in [
    "Open networking.",
    "Industry–academia interactions.",
    "HR and placement discussions.",
    "Collaboration opportunities.",
    "IEEE membership engagement.",
]:
    bullet(item)
body("Objective: to facilitate meaningful conversations and long-term professional connections among all "
     "stakeholders.", italic=True)

h2("Closing & Lunch — 12:50 PM to 1:00 PM")
for item in [
    "Closing Remarks.",
    "Acknowledgements.",
    "Group Announcements.",
]:
    bullet(item)
body("1:00 PM onwards — Lunch & informal networking for leaders, professionals, IEEE ExCom members and volunteers.")

# ---------------- 8. POST-LUNCH ----------------
h1("8.  Post-Lunch — IEEE Computer Society 80 Years Celebration")
body("Following lunch, the venue and stage transition to the IEEE Computer Society 80 Years Celebration programme.")

# ---------------- 9. REGISTRATION STRATEGY ----------------
h1("9.  Registration Strategy — Two Separate Tracks")
make_table(
    ["Form", "For"],
    [
        ["Form 1 — Leadership Round Table", "Invited leadership participants"],
        ["Form 2 — CS Professional Connect", "Professionals · TPOs · Placement Coordinators · Academicians"],
    ],
    widths=[2.6, 3.9],
)
body("Separate invitation designs and communication plans will be created for each audience.", italic=True)

h2("Live registration links")
LEADERS_FORM = "https://docs.google.com/forms/d/e/1FAIpQLSexJ1lUZZT_RarcPOsKUfdkxXlZBlsh_V__29rsmI6FnUFO_A/viewform"
PROS_FORM = "https://docs.google.com/forms/d/e/1FAIpQLSd686aQcAj5rE7ZykW77mTAeotlMuyp3w0Cfn-kW6d8lr_yZw/viewform"
p = doc.add_paragraph()
r = p.add_run("Form 1 · Leadership Round Table — ")
r.bold = True
r.font.size = Pt(10.5)
add_hyperlink(p, LEADERS_FORM, LEADERS_FORM)
p.paragraph_format.space_after = Pt(2)
p = doc.add_paragraph()
r = p.add_run("Form 2 · CS Professional Connect — ")
r.bold = True
r.font.size = Pt(10.5)
add_hyperlink(p, PROS_FORM, PROS_FORM)
p.paragraph_format.space_after = Pt(2)

# ---------------- 10. EVENT OPERATIONS ----------------
h1("10.  Event Operations")
h2("Registration Desk")
for item in ["Check-in.", "ID card distribution.", "Event information support."]:
    bullet(item)
h2("IEEE Membership Desk")
for item in ["Membership information.", "On-the-spot registration support.", "IEEE benefits awareness."]:
    bullet(item)
h2("Membership Incentives")
bullet("Selected early registrants may receive IEEE membership fee waivers.")
bullet("Participants completing IEEE membership registration during the event will receive special IEEE goodies.")

# ---------------- 11. IDENTITY CARDS ----------------
h1("11.  Identity Cards")
body("Separate colour-coded badges for each group:")
for item in ["Leaders", "Professionals", "IEEE ExCom", "Volunteers", "Organising Team"]:
    bullet(item)

# ---------------- 12. SECURITY & LOGISTICS ----------------
h1("12.  Security & Logistics")
for item in [
    "Dedicated storage room for goodies.",
    "Restricted access.",
    "Trusted volunteer team.",
    "Inventory tracking before and after distribution.",
    "Goodie management and allocation records.",
]:
    bullet(item)

# ---------------- 13. WEBSITE & INVITATION ----------------
h1("13.  Website & Invitation Experience")
body("The event website should reflect:")
for item in [
    "Premium leadership forum positioning.",
    "Modern and elegant design.",
    "Strong AI and industry focus.",
    "Clear event flow.",
    "Seamless registration experience.",
]:
    bullet(item)
h2("Leader Invitations")
for item in [
    "Premium and professional.",
    "Exclusive positioning.",
    "Visually distinct from professional invitations.",
]:
    bullet(item)
h2("Professional Invitations")
for item in [
    "Focused on networking opportunities.",
    "Industry–academia collaboration.",
    "IEEE engagement and community building.",
]:
    bullet(item)

# ---------------- 14. SUGGESTED LEADERS ----------------
h1("14.  Suggested Leaders")
body("Working list for the leadership invitations (to be finalised):", italic=True, color=GREY)
make_table(
    ["Name", "Affiliation / Note"],
    [
        ["Jayshree", "—"],
        ["Jayanth", "—"],
        ["Nithin Mohan", "CEO, Cambridge Group"],
        ["Sabrina", "—"],
        ["Mohin", "—"],
        ["Abhishek", "—"],
        ["Jyothi", "—"],
        ["RV College Chairperson", "RV College"],
        ["Kishor", "TCIEM"],
        ["Changappa", "Bosch"],
        ["Veena Rao", "—"],
        ["Rajesh", "—"],
        ["Vineesh", "—"],
        ["Prashanth", "IEEE"],
        ["Deepa", "IEEE CS ExCom"],
        ["Ashwini", "—"],
        ["Sujata", "—"],
        ["Sanil", "BMSIT"],
        ["Yashwanth", "—"],
        ["Shobha Ma'am", "—"],
        ["Vidya Kabra", "—"],
        ["Himanshu", "—"],
    ],
    widths=[3.0, 3.5],
)

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("IEEE CS Pro 2026  ·  IEEE Computer Society – Bangalore Chapter  ·  ieeecsbangalore.org")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = GREY

out = r"D:\projects\claude\CSPro\IEEE_CS_Pro_2026_Event_Plan.docx"
doc.save(out)
print("Saved:", out)

# Keep the user's working copy on the Desktop in sync
desktop = r"C:\Users\imamf\OneDrive\Desktop\IEEE_CS_Pro_2026_Event_Plan.docx"
try:
    shutil.copyfile(out, desktop)
    print("Copied to:", desktop)
except OSError as e:
    print("Could not copy to Desktop:", e)
